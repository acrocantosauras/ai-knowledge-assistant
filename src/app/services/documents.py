from __future__ import annotations

import hashlib
from pathlib import Path
from uuid import UUID

from fastapi import HTTPException, UploadFile, status
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import get_settings
from app.db.models.document import Document
from app.db.models.document_chunk import DocumentChunk
from app.db.models.user import User
from app.services.rag import (
    ChunkingConfig,
    create_document_chunks,
    create_embeddings_for_chunks,
    get_embedding_service,
)

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".docx",
}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024


def _normalize_filename(filename: str) -> str:
    return Path(filename).name or "upload"


def _detect_content_type(filename: str, declared_type: str | None) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return "application/pdf"
    if lower_name.endswith(".txt"):
        return "text/plain"
    if lower_name.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return declared_type or "application/octet-stream"


def _validate_file(file: UploadFile) -> None:
    filename = _normalize_filename(file.filename or "")
    extension = Path(filename).suffix.lower()
    content_type = _detect_content_type(filename, file.content_type)

    is_supported_ext = extension in SUPPORTED_EXTENSIONS
    is_supported_mime = content_type in SUPPORTED_MIME_TYPES
    if not is_supported_ext and not is_supported_mime:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported file type. Supported: PDF, TXT, and DOCX.",
        )

    if file.size is not None and file.size > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File too large. Maximum size is 10 MB.",
        )


def _read_text_from_upload(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return file_bytes.decode("utf-8")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF support is unavailable in this environment.",
            ) from exc

        reader = PdfReader(bytes(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOCX support is unavailable in this environment.",
            ) from exc

        doc = DocxDocument(file_bytes)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Unsupported file type. Supported: PDF, TXT, and DOCX.",
    )


def extract_pdf_metadata(file_bytes: bytes) -> dict:
    """Extract metadata from PDF file."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return {}

    try:
        reader = PdfReader(bytes(file_bytes))
        meta = reader.metadata
        if not meta:
            return {}

        return {
            "author": meta.get("/Author", ""),
            "creator": meta.get("/Creator", ""),
            "producer": meta.get("/Producer", ""),
            "subject": meta.get("/Subject", ""),
            "title": meta.get("/Title", ""),
            "creation_date": meta.get("/CreationDate", ""),
            "modification_date": meta.get("/ModDate", ""),
            "page_count": len(reader.pages),
        }
    except Exception:
        return {}


def extract_docx_metadata(file_bytes: bytes) -> dict:
    """Extract metadata from DOCX file."""
    try:
        from docx import Document as DocxDocument
        from docx.opc.coreprops import CoreProperties
    except ImportError:
        return {}

    try:
        doc = DocxDocument(file_bytes)
        props: CoreProperties = doc.core_properties

        return {
            "author": props.author or "",
            "category": props.category or "",
            "comments": props.comments or "",
            "content_status": props.content_status or "",
            "created": props.created.isoformat() if props.created else "",
            "identifier": props.identifier or "",
            "keywords": props.keywords or "",
            "language": props.language or "",
            "last_modified_by": props.last_modified_by or "",
            "modified": props.modified.isoformat() if props.modified else "",
            "revision": props.revision or "",
            "subject": props.subject or "",
            "title": props.title or "",
            "version": props.version or "",
            "paragraph_count": len(doc.paragraphs),
        }
    except Exception:
        return {}


def extract_text_metadata(text: str) -> dict:
    """Extract basic metadata from text content."""
    lines = text.splitlines()
    words = text.split()
    return {
        "line_count": len(lines),
        "word_count": len(words),
        "char_count": len(text),
    }


async def upload_document(
    session: AsyncSession,
    user: User,
    file: UploadFile,
) -> Document:
    if file.filename is None or not file.filename.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="A file is required.",
        )

    _validate_file(file)

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty.",
        )

    text = _read_text_from_upload(file_bytes, file.filename)
    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Uploaded file does not contain readable text.",
        )

    checksum = hashlib.sha256(file_bytes).hexdigest()
    excerpt = text[:200].strip().replace("\n", " ")

    # Extract file-specific metadata
    filename = _normalize_filename(file.filename)
    content_type = _detect_content_type(file.filename, file.content_type)
    suffix = Path(filename).suffix.lower()

    file_metadata = {
        "filename": filename,
        "content_length": len(file_bytes),
        "mime_type": content_type,
        "extracted_chars": len(text),
        "excerpt": excerpt,
    }

    # Add file-type specific metadata
    if suffix == ".pdf":
        file_metadata["pdf_metadata"] = extract_pdf_metadata(file_bytes)
    elif suffix == ".docx":
        file_metadata["docx_metadata"] = extract_docx_metadata(file_bytes)
    elif suffix == ".txt":
        file_metadata["text_metadata"] = extract_text_metadata(text)

    doc = Document(
        user_id=user.id,
        title=filename,
        source_type="upload",
        source_uri=None,
        content_type=content_type,
        status="processed",
        checksum=checksum,
        content=text,
        document_metadata=file_metadata,
    )
    session.add(doc)
    await session.flush()

    # Create chunks from the extracted text
    settings = get_settings()
    chunk_config = ChunkingConfig(
        chunk_size=settings.chunk_size,
        overlap=settings.chunk_overlap,
    )
    chunks = await create_document_chunks(session, doc, text, chunk_config)

    # Generate embeddings for chunks
    embedding_service = get_embedding_service()
    await create_embeddings_for_chunks(session, chunks, embedding_service)

    await session.commit()
    await session.refresh(doc)
    return doc


async def get_document_by_id(
    session: AsyncSession,
    document_id: UUID,
    user_id: UUID,
) -> Document | None:
    result = await session.execute(
        select(Document).where(Document.id == document_id, Document.user_id == user_id)
    )
    return result.scalar_one_or_none()


async def list_user_documents(
    session: AsyncSession, user_id: UUID, limit: int = 50, offset: int = 0
) -> tuple[list[Document], int]:
    """List documents belonging to the specified user."""
    query = (
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    documents = list((await session.execute(query)).scalars())

    # Count all user documents
    count_query = select(func.count(Document.id)).where(Document.user_id == user_id)
    total = (await session.execute(count_query)).scalar()

    return documents, total


async def get_document(
    session: AsyncSession, user_id: UUID, document_id: UUID
) -> Document | None:
    """Fetch the document object if it exists."""
    query = (
        select(Document)
        .where(Document.id == document_id, Document.user_id == user_id)
        .options(selectinload(Document.chunks))
    )
    return (await session.execute(query)).scalar()


async def delete_document(
    session: AsyncSession, user_id: UUID, document_id: UUID
) -> bool:
    """Delete a document and its associated chunks."""
    doc = await get_document(session, user_id, document_id)

    if not doc:
        return False  # No document to delete

    chunks_query = select(DocumentChunk).where(DocumentChunk.document_id == document_id)
    chunks = list((await session.execute(chunks_query)).scalars())
    for chunk in chunks:
        await session.delete(chunk)

    await session.delete(doc)
    await session.commit()
    return True
